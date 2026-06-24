# -*- coding: utf-8 -*-
"""
Regenerates Plan_de_Respuestas.docx with the format:
**Pregunta N:** [question text]
**Respuesta:** [answer text]

For the course 'Técnicas de Formulación de Proyectos'.
"""

import sys
import os
from datetime import date

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from apa7 import APA7Doc
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

OUTPUT_DIR = r"{{USER_DIR}}\OneDrive\Documentos\Archivos OpenCode\Documentos"
TODAY = date.today().strftime("%d de %B de %Y")
STUDENT = "[Nombre del Estudiante]"
UNIVERSITY = "[Nombre de la Universidad]"
COURSE = "Técnicas de Formulación de Proyectos"
INSTRUCTOR = "[Nombre del Profesor]"


def add_bold_normal_paragraph(doc, bold_text, normal_text):
    """Add a paragraph with a bold segment followed by normal text, APA7 indented."""
    p = doc.doc.add_paragraph()
    # Bold run
    run_bold = p.add_run(bold_text)
    run_bold.font.name = APA7Doc.FONT
    run_bold.font.size = APA7Doc.FONT_SIZE
    run_bold.bold = True
    # Normal run
    run_normal = p.add_run(normal_text)
    run_normal.font.name = APA7Doc.FONT
    run_normal.font.size = APA7Doc.FONT_SIZE
    # Paragraph formatting (APA7 style)
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = APA7Doc.LINE_SPACING
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = APA7Doc.INDENT
    return p


def add_qa_pair(doc, question_num, question_text, answer_text):
    """Add a Q&A pair: bold question label + text, then bold Respuesta + answer."""
    add_bold_normal_paragraph(doc, f"Pregunta {question_num}: ", question_text)
    add_bold_normal_paragraph(doc, "Respuesta: ", answer_text)


def generar_plan_respuestas():
    doc = APA7Doc()
    doc.add_title_page(
        "Plan de Respuestas — Técnicas de Formulación de Proyectos",
        STUDENT,
        UNIVERSITY,
        course=COURSE,
        instructor=INSTRUCTOR,
        date=TODAY,
    )
    doc.add_toc_page()

    # ========================================================================
    # SECTION 1 — Respuestas a Preguntas de 20 Puntos
    # ========================================================================
    doc.add_heading_level1("Respuestas a Preguntas de 20 Puntos")

    add_qa_pair(doc, 1,
        "De acuerdo con el Project Management Institute (PMI), ¿qué es un proyecto y cuáles son sus características fundamentales?",
        "Un proyecto es un esfuerzo temporal que se lleva a cabo para crear un producto, servicio o resultado único (Project Management Institute, 2021). Sus características fundamentales son: (a) temporalidad — tiene un inicio y un fin definidos; (b) unicidad — el resultado es diferente a cualquier otro producto o servicio; y (c) elaboración progresiva — se desarrolla en pasos sucesivos, aumentando el detalle a medida que avanza."
    )

    add_qa_pair(doc, 2,
        "Explique la diferencia fundamental entre un proyecto y las operaciones dentro de una organización.",
        "La diferencia fundamental radica en dos aspectos: (a) Temporalidad: los proyectos son temporales con fechas de inicio y fin definidas, mientras que las operaciones son continuas y repetitivas, manteniendo la organización en funcionamiento. (b) Unicidad: los proyectos producen entregables únicos, mientras que las operaciones implican la ejecución repetitiva de tareas estandarizadas. Mientras las operaciones sostienen la organización, los proyectos permiten el cambio y la mejora."
    )

    add_qa_pair(doc, 3,
        "¿Cuáles son las cuatro fases del ciclo de vida de un proyecto según la guía del PMBOK?",
        "Las cuatro fases son: (1) Inicio — se define el proyecto, se identifican los interesados y se obtiene la autorización; (2) Planificación — se establece el alcance, los objetivos y el curso de acción; (3) Ejecución — se completa el trabajo definido en el plan de gestión; (4) Cierre — se formaliza la aceptación del producto, servicio o resultado y se liberan los recursos del proyecto."
    )

    add_qa_pair(doc, 4,
        "¿Qué es una Oficina de Gestión de Proyectos (PMO) y cuáles son sus funciones principales?",
        "Una PMO es una unidad organizacional centralizada que supervisa la gestión de proyectos dentro de una organización. Entre sus funciones principales se encuentran: estandarizar los procesos de gestión de proyectos, proporcionar metodologías y herramientas, gestionar recursos compartidos entre proyectos, facilitar la capacitación del personal, monitorear el rendimiento del portafolio de proyectos y asegurar la alineación con los objetivos estratégicos de la organización."
    )

    add_qa_pair(doc, 5,
        "Defina el modelo de las Cinco Fuerzas de Porter y enumérelas.",
        "Las Cinco Fuerzas de Porter son un marco para analizar la competitividad y el atractivo de una industria. Estas fuerzas son: (1) Amenaza de nuevos competidores o entrantes; (2) Poder de negociación de los compradores o clientes; (3) Poder de negociación de los proveedores; (4) Amenaza de productos o servicios sustitutos; (5) Rivalidad entre los competidores existentes. El análisis de estas fuerzas permite determinar el potencial de rentabilidad de un mercado."
    )

    # ========================================================================
    # SECTION 2 — Respuestas a Preguntas de 30 Puntos
    # ========================================================================
    doc.add_heading_level1("Respuestas a Preguntas de 30 Puntos")

    add_qa_pair(doc, 6,
        "Compare los criterios de éxito tradicionales de un proyecto (presupuesto, tiempo, alcance) con los criterios ampliados propuestos por el PMI y Hughes. ¿En qué se diferencian?",
        'Los criterios tradicionales de éxito (el "Triángulo de Hierro") miden el éxito de un proyecto mediante tres restricciones: finalizar dentro del presupuesto, dentro del plazo establecido y cumpliendo con el alcance definido. El PMI (2021) amplía estos criterios para incluir la satisfacción del cliente, la gestión de las expectativas de los interesados y la calidad de los entregables. Hughes añade dimensiones adicionales como el éxito del negocio a largo plazo, los beneficios organizacionales y el logro de objetivos estratégicos. La gestión moderna reconoce que un proyecto puede cumplir con el triple constraint pero fracasar si no genera valor para los interesados; inversamente, un proyecto que excede el presupuesto pero genera un valor estratégico excepcional puede considerarse exitoso.'
    )

    add_qa_pair(doc, 7,
        "Explique cómo la Matriz BCG (Boston Consulting Group) ayuda en la toma de decisiones de posicionamiento de mercado para un proyecto.",
        "La Matriz BCG clasifica productos o unidades de negocio en cuatro cuadrantes según la tasa de crecimiento del mercado y la participación de mercado relativa: (1) Estrellas — alto crecimiento y alta participación, requieren inversión para mantener el liderazgo; (2) Interrogantes — alto crecimiento pero baja participación, requieren análisis para decidir entre invertir o desinvertir; (3) Vacas Lecheras — bajo crecimiento y alta participación, generan efectivo con inversión mínima; (4) Perros — bajo crecimiento y baja participación, candidatos para desinversión. En la formulación de proyectos, la Matriz BCG ayuda a decidir la asignación de recursos: las Estrellas deben recibir inversión, los Interrogantes necesitan estudios de mercado para determinar su viabilidad, las Vacas Lecheras financian otras iniciativas, y los Perros deben minimizarse."
    )

    add_qa_pair(doc, 8,
        "Explique la diferencia entre el Valor Actual Neto (VAN) y la Tasa Interna de Retorno (TIR) en los estudios de factibilidad económica. ¿Cuál es la regla de decisión para cada uno?",
        "El VAN (Valor Actual Neto) calcula el valor presente de los flujos de efectivo futuros menos la inversión inicial. Un VAN positivo indica que el proyecto genera valor por encima de la tasa de descuento. Se expresa en unidades monetarias (por ejemplo, dólares) y proporciona una medida absoluta de valor. La TIR (Tasa Interna de Retorno) es la tasa de descuento que hace que el VAN sea igual a cero, representando la tasa de rendimiento esperada del proyecto en términos porcentuales. La regla de decisión para el VAN es: aceptar el proyecto si VAN > 0. Para la TIR: aceptar el proyecto si TIR > tasa de descuento (costo de capital). La diferencia clave es que el VAN proporciona una medida absoluta de valor, mientras que la TIR ofrece un rendimiento porcentual relativo. Para proyectos mutuamente excluyentes, se prefiere el VAN como criterio de decisión."
    )

    add_qa_pair(doc, 9,
        "Describa cómo se utiliza la cadena de valor para realizar un análisis interno en la formulación de proyectos. Identifique las actividades primarias y de apoyo.",
        "El modelo de la cadena de valor divide las actividades organizacionales en actividades primarias y de apoyo. Las actividades primarias incluyen: logística interna, operaciones, logística externa, marketing y ventas, y servicio postventa. Las actividades de apoyo incluyen: adquisiciones (compras), gestión de recursos humanos, desarrollo tecnológico e infraestructura organizacional. En la formulación de proyectos, la cadena de valor se utiliza para identificar ventajas competitivas y oportunidades de mejora mediante el análisis del costo y el potencial de diferenciación de cada actividad. El análisis revela qué actividades agregan valor al producto final y ayuda al equipo del proyecto a diseñar iniciativas que optimicen estas actividades, reduzcan costos o mejoren la diferenciación."
    )

    add_qa_pair(doc, 10,
        "Compare los diferentes niveles del marco legal que afectan la factibilidad de un proyecto, desde los tratados internacionales hasta las normas internas. Explique la jerarquía.",
        "La jerarquía del marco legal incluye: (1) Tratados Internacionales — acuerdos vinculantes entre naciones que en muchas jurisdicciones prevalecen sobre el derecho interno; (2) Constitución — la ley suprema del país, que establece los principios fundamentales; (3) Leyes — emitidas por el poder legislativo, regulan áreas específicas como derecho laboral, tributario o ambiental; (4) Reglamentos (Decretos) — normas del poder ejecutivo que implementan las leyes; (5) Políticas — directrices administrativas emitidas por agencias gubernamentales; (6) Normas Internas — reglas y procedimientos específicos de una organización. Cada nivel condiciona o habilita el proyecto de manera diferente. Un proyecto debe cumplir con todos los niveles aplicables; el incumplimiento en niveles superiores (por ejemplo, constitucional) invalida las autorizaciones de niveles inferiores."
    )

    # ========================================================================
    # SECTION 3 — Respuestas a Preguntas de 50 Puntos
    # ========================================================================
    doc.add_heading_level1("Respuestas a Preguntas de 50 Puntos")

    add_qa_pair(doc, 11,
        (
            "Usted está gestionando un proyecto de construcción residencial que actualmente está 3 meses retrasado y 15% "
            "por encima del presupuesto. Sin embargo, el cliente ha expresado satisfacción con la calidad de los materiales "
            "y la mano de obra. Analice esta situación utilizando los criterios de éxito discutidos en clase (presupuesto, "
            "tiempo, satisfacción del cliente, criterios de Hughes, procesos del PMI). ¿Qué dimensiones están teniendo éxito "
            "y cuáles están fallando? ¿Qué acciones correctivas recomendaría?"
        ),
        (
            "Este caso presenta un conflicto entre las dimensiones tradicionales del éxito y los criterios ampliados. "
            "Dimensiones que están fallando: (a) Tiempo: 3 meses de retraso indica una desviación significativa del "
            "cronograma, lo que incumple el criterio tradicional del Triángulo de Hierro; (b) Presupuesto: 15% de "
            "sobrecosto supera la tolerancia típica de gestión de proyectos, lo que indica deficiencias en la "
            "planificación de costos o en el control durante la ejecución. "
            "Dimensiones que están teniendo éxito: (a) Satisfacción del cliente: el cliente está satisfecho con la "
            "calidad, cumpliendo con un criterio clave del PMI; (b) Calidad: según los criterios de Hughes, la calidad "
            "del producto final es un indicador de éxito a largo plazo. "
            "Acciones correctivas recomendadas: (1) Realizar una gestión del valor ganado (EVM) para cuantificar la "
            "desviación de cronograma y costo; (2) Implementar un plan de recuperación con cronograma comprimido (fast "
            "tracking o crashing); (3) Revisar el presupuesto restante y reasignar recursos; (4) Establecer hitos de "
            "control más frecuentes; (5) Documentar las lecciones aprendidas para evitar desviaciones similares; "
            "(6) Evaluar si el sobrecosto puede compensarse con el alto nivel de satisfacción del cliente, que podría "
            "traducirse en negocios futuros o referencias."
        )
    )

    add_qa_pair(doc, 12,
        (
            "Una empresa de tecnología desea lanzar un innovador dispositivo para hogar inteligente en un mercado "
            "actualmente dominado por tres competidores bien establecidos. Utilizando el modelo de las Cinco Fuerzas de "
            "Porter, evalúe la viabilidad de ingresar a este mercado. Para cada fuerza, identifique factores específicos "
            "que afectan el atractivo del mercado y proponga una estrategia para mitigar los riesgos identificados."
        ),
        (
            "1. Amenaza de nuevos entrantes: Alta si existen bajas barreras de entrada tecnológicas. Estrategia: "
            "Solicitar patentes para la tecnología innovadora y establecer acuerdos de distribución exclusivos. "
            "2. Poder de negociación de compradores: Alto si los costos de cambio son bajos (típico en electrónica de "
            "consumo). Estrategia: Diferenciación mediante características únicas de software e integración con "
            "ecosistemas existentes. "
            "3. Poder de negociación de proveedores: Medio-alto si los componentes son especializados. Estrategia: "
            "Establecer contratos a largo plazo y considerar integración vertical parcial para componentes críticos. "
            "4. Amenaza de sustitutos: Alta, existen múltiples alternativas para el control del hogar. Estrategia: "
            "Enfocarse en la experiencia de usuario superior y la interoperabilidad. "
            "5. Rivalidad entre competidores existentes: Muy alta con tres actores establecidos. Estrategia: Identificar "
            "un nicho desatendido (por ejemplo, seguridad para adultos mayores o integración con energías renovables) "
            "para evitar la competencia frontal. Conclusión: El mercado es de atractivo moderado a bajo; la viabilidad "
            "depende de una estrategia de nicho clara y una diferenciación significativa."
        )
    )

    add_qa_pair(doc, 13,
        (
            "Una bebida orgánica recién lanzada ha sido clasificada en un mercado con una tasa de crecimiento anual "
            "del 15%, pero actualmente posee solo el 5% de participación de mercado. Utilizando la Matriz BCG, clasifique "
            "este producto y analice sus implicaciones estratégicas. ¿Debería la empresa invertir fuertemente, mantener "
            "el gasto actual o desinvertir? Justifique su recomendación considerando flujo de caja, estrategia "
            "competitiva y dinámica del mercado."
        ),
        (
            "La bebida orgánica se clasifica como Interrogante (alto crecimiento del mercado al 15%, baja participación "
            "del 5%). Implicaciones estratégicas: (a) Requiere inversión significativa para aumentar la participación de "
            "mercado y convertirse en estrella; (b) Consume efectivo en lugar de generarlo; (c) El mercado de alto "
            "crecimiento indica demanda insatisfecha. Recomendación: Invertir selectivamente si se identifican las "
            "barreras específicas que limitan la participación (distribución, conocimiento de marca, precio). "
            "Estrategias concretas: (1) Realizar un estudio de mercado específico para entender por qué la participación "
            "es baja en un mercado de alto crecimiento; (2) Si la barrera es distribución, buscar alianzas estratégicas "
            "con cadenas de retail; (3) Campaña de marketing agresiva dirigida al segmento objetivo; (4) Evaluar si la "
            "capacidad de producción puede escalar. Si después de 12-18 meses no hay mejora en la participación, "
            "considerar desinversión. No se recomienda gasto corriente sin inversión, porque el producto Interrogante "
            "debe migrar a Estrella o desinvertirse para no consumir recursos indefinidamente."
        )
    )

    add_qa_pair(doc, 14,
        (
            "Un proyecto de expansión manufacturera presenta los siguientes datos: Inversión inicial: $500,000. "
            "Flujos de efectivo proyectados: Año 1: $100,000; Año 2: $150,000; Año 3: $200,000; Año 4: $250,000; "
            "Año 5: $300,000. La tasa de descuento de la empresa es del 8%. Calcule e interprete el VAN. Si la TIR "
            "se calcula en 18.5%, ¿qué nos dice esto sobre la viabilidad del proyecto? Analice los riesgos asociados "
            "con estas estimaciones y recomiende análisis de sensibilidad."
        ),
        (
            "Cálculo del VAN: VAN = -$500,000 + $100,000/(1.08)¹ + $150,000/(1.08)² + $200,000/(1.08)³ + "
            "$250,000/(1.08)⁴ + $300,000/(1.08)⁵ = -$500,000 + $92,592.59 + $128,600.82 + $158,766.45 + "
            "$183,756.44 + $204,174.32 = $267,890.62. "
            "Interpretación: VAN > $0, el proyecto genera valor neto positivo. Por cada dólar invertido, se recupera "
            "el capital y se generan $267,890.62 adicionales en valor presente. "
            "TIR: La TIR del 18.5% es mayor que la tasa de descuento del 8%, lo que confirma la viabilidad del proyecto "
            "desde ambos criterios. La TIR indica que el proyecto rinde 18.5% anual, superando ampliamente el costo de "
            "capital. "
            "Riesgos: (1) Los flujos de efectivo son estimaciones, sujetas a variaciones en ventas, costos de producción "
            "y condiciones económicas; (2) La tasa de descuento del 8% puede cambiar si aumentan las tasas de interés. "
            "Análisis de sensibilidad recomendado: (1) Calcular el VAN con escenarios optimista (crecimiento del 10% en "
            "flujos), pesimista (caída del 20% en flujos) y base; (2) Determinar el punto de equilibrio — el flujo "
            "mínimo del Año 5 que mantiene VAN > $0; (3) Calcular la TIR modificada (TIRM) para corregir el supuesto "
            "de reinversión implícito en la TIR tradicional."
        )
    )

    add_qa_pair(doc, 15,
        (
            "Una universidad pública propone implementar un nuevo programa de licenciatura completamente en línea. "
            "La misión institucional establece \"educación superior accesible\" y la visión es \"ser líder regional "
            "en innovación educativa\". Utilizando el marco de alineación estratégica, formule una justificación "
            "integral para la autorización de este proyecto. Su análisis debe incluir: alineación misión-visión, "
            "perspectivas del Balanced Scorecard, análisis interno de cadena de valor, análisis competitivo externo, "
            "y la necesidad específica de mercado u organizacional que justifica este proyecto."
        ),
        (
            "Alineación Misión-Visión: El programa en línea se alinea directamente con la misión de \"educación superior "
            "accesible\" al eliminar barreras geográficas y de horario; también contribuye a la visión de \"ser líder "
            "regional en innovación educativa\" al adoptar un modelo de entrega moderno y tecnológico. "
            "Balanced Scorecard: (a) Perspectiva Financiera — el programa genera ingresos por matrícula con costos "
            "operativos potencialmente menores; (b) Perspectiva del Cliente (Estudiantes) — satisface la demanda de "
            "flexibilidad educativa; (c) Perspectiva de Procesos Internos — requiere el diseño de procesos de "
            "enseñanza-aprendizaje virtuales eficientes; (d) Perspectiva de Aprendizaje y Crecimiento — la universidad "
            "desarrolla capacidades en educación digital. "
            "Cadena de Valor (Análisis Interno): Las actividades primarias incluyen diseño curricular, plataforma "
            "tecnológica, admisión virtual, y soporte al estudiante. Las actividades de apoyo incluyen desarrollo "
            "profesoral en competencias digitales, infraestructura tecnológica, y gestión administrativa. "
            "Análisis Externo: (a) Demanda de mercado — creciente preferencia por la educación en línea post-pandemia; "
            "(b) Cinco Fuerzas — competencia de otras universidades con programas en línea, poder de negociación de "
            "estudiantes que comparan opciones, amenaza de plataformas educativas globales como Coursera o edX. "
            "Autorización del Proyecto: Se justifica por necesidad organizacional (ampliar alcance institucional), "
            "solicitud del cliente (demanda estudiantil por modalidad virtual), y avance tecnológico (madurez de "
            "plataformas de aprendizaje en línea). Se recomienda un estudio de factibilidad que incluya análisis de "
            "mercado estudiantil, viabilidad técnica de la plataforma, y proyección financiera a 5 años."
        )
    )

    # ========================================================================
    # SECTION 4 — Preguntas Adicionales y Respuestas
    # ========================================================================
    doc.add_heading_level1("Preguntas Adicionales y Respuestas")

    add_qa_pair(doc, 16,
        "¿Qué es el Balanced Scorecard (BSC) y cómo se relaciona con la formulación de proyectos?",
        "El BSC es una herramienta de gestión estratégica que traduce la misión y visión de una organización en objetivos medibles desde cuatro perspectivas: financiera, del cliente, de procesos internos, y de aprendizaje y crecimiento. En la formulación de proyectos, el BSC asegura que los proyectos estén alineados con los objetivos estratégicos de la organización, proporcionando un marco para seleccionar y priorizar proyectos que contribuyan a metas organizacionales medibles."
    )

    add_qa_pair(doc, 17,
        "¿Qué es el benchmarking y cómo se utiliza en un estudio de mercado?",
        "El benchmarking es un proceso sistemático de comparación de productos, servicios o procesos con los de organizaciones líderes o competidores directos para identificar mejores prácticas y áreas de mejora. En un estudio de mercado para un proyecto, el benchmarking permite: identificar estándares de la industria, analizar las estrategias de la competencia, determinar precios de referencia, y establecer objetivos de desempeño realistas basados en el mercado."
    )

    add_qa_pair(doc, 18,
        "¿Qué es la elasticidad de la demanda y por qué es importante para un estudio de factibilidad comercial?",
        "La elasticidad de la demanda mide la sensibilidad de la cantidad demandada ante cambios en el precio. Es importante porque: (a) determina cómo las variaciones de precio afectarán los ingresos del proyecto; (b) ayuda a fijar precios óptimos; (c) indica el poder de mercado del proyecto. Una demanda elástica significa que pequeños cambios en el precio causan grandes cambios en la cantidad demandada; una demanda inelástica significa lo contrario."
    )

    add_qa_pair(doc, 19,
        "¿Cuáles son las cinco fuentes de autorización de proyectos según el PMI?",
        "Las cinco fuentes son: (1) Demanda del mercado — un proyecto surge para satisfacer una necesidad del mercado; (2) Necesidad organizacional — cambios internos requieren un proyecto; (3) Solicitud del cliente — un cliente externo solicita un producto o servicio; (4) Avance tecnológico — la adopción de nueva tecnología requiere un proyecto; (5) Requisito legal — una nueva ley o regulación exige cambios en la organización."
    )

    add_qa_pair(doc, 20,
        "¿Cuál es la diferencia entre normas obligatorias y mejores prácticas en la gestión de proyectos?",
        'Las normas obligatorias (como leyes y regulaciones) son de cumplimiento forzoso y su incumplimiento puede conllevar sanciones legales o regulatorias. Las mejores prácticas y estándares (como la guía del PMBOK o las normas ISO) no son obligatorias pero proporcionan marcos de referencia y metodologías probadas que ofrecen ventajas competitivas. Adoptar mejores prácticas es voluntario pero recomendado para mejorar la eficiencia y eficacia de la gestión de proyectos.'
    )

    add_qa_pair(doc, 21,
        "Explique la relación entre la misión, la visión y los objetivos estratégicos de una organización.",
        "La misión define el propósito actual de la organización (qué hace, para quién y cómo). La visión describe el estado futuro deseado (dónde quiere estar en el largo plazo). Los objetivos estratégicos son metas específicas y medibles que traducen la visión en acciones concretas, conectando el presente (misión) con el futuro (visión). Los proyectos se derivan de estos objetivos estratégicos, asegurando que cada iniciativa contribuya al logro de la visión organizacional."
    )

    add_qa_pair(doc, 22,
        "¿Qué es un estudio de factibilidad y cuáles son sus componentes principales?",
        "Un estudio de factibilidad es un análisis integral que determina la viabilidad de un proyecto antes de su ejecución. Sus componentes principales son: (1) Estudio de Mercado (Factibilidad Comercial) — análisis de demanda, oferta, precios y competencia; (2) Estudio Técnico — análisis de recursos técnicos, tecnología, procesos y localización; (3) Estudio Legal — análisis del marco regulatorio aplicable; (4) Estudio Económico-Financiero — evaluación de VAN, TIR, período de recuperación y análisis de sensibilidad."
    )

    add_qa_pair(doc, 23,
        "Explique la relación entre el VAN y la TIR como criterios de evaluación de inversiones.",
        "El VAN y la TIR son complementarios: el VAN mide el valor absoluto generado por el proyecto en términos monetarios actuales, mientras que la TIR mide la rentabilidad relativa en términos porcentuales. La regla de decisión es consistente: si VAN > 0, entonces TIR > tasa de descuento (para proyectos convencionales). Sin embargo, pueden surgir conflictos al comparar proyectos mutuamente excluyentes de diferente escala o duración. En tales casos, el VAN es el criterio superior porque maximiza el valor absoluto para la organización."
    )

    add_qa_pair(doc, 24,
        "¿Qué factores deben considerarse en un estudio de factibilidad técnica?",
        "Los factores clave incluyen: (1) Características técnicas del producto o servicio; (2) Requerimientos de infraestructura (terrenos, edificios, instalaciones); (3) Requerimientos de maquinaria y equipo; (4) Proceso de producción u operación; (5) Disponibilidad de insumos y materias primas; (6) Requerimientos de personal técnico; (7) Localización y distribución geográfica; (8) Impacto ambiental y medidas de mitigación; (9) Plan de mantenimiento y soporte técnico; (10) Cronograma de implementación técnica."
    )

    add_qa_pair(doc, 25,
        "¿Cómo se relaciona el estudio de mercado (comercial) con el estudio técnico en la formulación de proyectos?",
        "El estudio de mercado determina la demanda esperada, el perfil del cliente, el precio y las estrategias de comercialización. El estudio técnico utiliza estos datos para dimensionar la capacidad de producción, seleccionar la tecnología adecuada, diseñar el proceso operativo y determinar la localización óptima. Por ejemplo, si el estudio de mercado proyecta una demanda de 10,000 unidades anuales, el estudio técnico debe diseñar una capacidad instalada que pueda satisfacer esta demanda. La relación es secuencial y recursiva: los resultados del estudio de mercado alimentan el estudio técnico, y las limitaciones técnicas pueden retroalimentar y ajustar las proyecciones de mercado."
    )

    # ========================================================================
    # Referencias
    # ========================================================================
    doc.add_heading_level1("Referencias")
    doc.add_reference(
        "Project Management Institute. (2021). Guía de los fundamentos para la dirección de proyectos "
        "(Guía del PMBOK) (7.ª ed.). Project Management Institute."
    )
    doc.add_reference(
        "Porter, M. E. (2008). Las cinco fuerzas competitivas que le dan forma a la estrategia. "
        "Harvard Business Review, 86(1), 78–93."
    )
    doc.add_reference(
        "Kaplan, R. S., & Norton, D. P. (1996). The balanced scorecard: Translating strategy into action. "
        "Harvard Business School Press."
    )
    doc.add_reference(
        "Henderson, B. D. (1970). The product portfolio. Boston Consulting Group."
    )

    filepath = os.path.join(OUTPUT_DIR, "Plan_de_Respuestas.docx")
    doc.save(filepath)
    print(f"[OK] Generado: {filepath}")


if __name__ == "__main__":
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    generar_plan_respuestas()
    print("[OK] Documento generado exitosamente.")

