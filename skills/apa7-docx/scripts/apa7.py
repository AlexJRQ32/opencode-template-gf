import docx
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
from datetime import date


class APA7Doc:
    FONT = "Times New Roman"
    FONT_SIZE = Pt(12)
    MARGIN = Cm(2.54)
    INDENT = Cm(1.27)
    LINE_SPACING = 1.5
    COLOR = RGBColor(0, 0, 0)

    def __init__(self):
        self.doc = docx.Document()
        self._set_default_font()
        self._set_margins()
        self._configure_styles()

    def _set_default_font(self):
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = self.FONT
        font.size = self.FONT_SIZE
        font.color.rgb = self.COLOR
        pf = style.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = self.LINE_SPACING
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _set_margins(self):
        for section in self.doc.sections:
            section.top_margin = self.MARGIN
            section.bottom_margin = self.MARGIN
            section.left_margin = self.MARGIN
            section.right_margin = self.MARGIN

    def _configure_styles(self):
        self._configure_heading_style("Heading 1", "1", True, False, WD_ALIGN_PARAGRAPH.CENTER)
        self._configure_heading_style("Heading 2", "2", True, False, WD_ALIGN_PARAGRAPH.LEFT)
        self._configure_heading_style("Heading 3", "3", True, True, WD_ALIGN_PARAGRAPH.LEFT)

    def _configure_heading_style(self, style_name, level, bold, italic, alignment):
        style = self.doc.styles[style_name]
        font = style.font
        font.name = self.FONT
        font.size = self.FONT_SIZE
        font.bold = bold
        font.italic = italic
        font.color.rgb = self.COLOR
        pf = style.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = self.LINE_SPACING
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.alignment = alignment
        pf.first_line_indent = Cm(0)
        self._set_outline_level(style, level)

    def _configure_heading_char_style(self, name, bold, italic):
        style = self.doc.styles[name]
        font = style.font
        font.name = self.FONT
        font.size = self.FONT_SIZE
        font.bold = bold
        font.italic = italic
        font.color.rgb = self.COLOR

    def _set_outline_level(self, style, level):
        pPr = style.element.get_or_add_pPr()
        outline = pPr.find(qn("w:outlineLvl"))
        if outline is None:
            outline = OxmlElement("w:outlineLvl")
            pPr.append(outline)
        outline.set(qn("w:val"), level)

    def _set_keep_with_next(self, paragraph):
        pPr = paragraph._p.get_or_add_pPr()
        keep = OxmlElement("w:keepNext")
        pPr.append(keep)

    def add_page_number(self, section=None):
        section = section or self.doc.sections[0]
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

        run = p.add_run()
        run.font.name = self.FONT
        run.font.size = self.FONT_SIZE
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_begin)

        run2 = p.add_run()
        run2.font.name = self.FONT
        run2.font.size = self.FONT_SIZE
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        run2._r.append(instr)

        run3 = p.add_run()
        run3.font.name = self.FONT
        run3.font.size = self.FONT_SIZE
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run3._r.append(fld_end)

    def add_toc_page(self, title="Índice"):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.name = self.FONT
        run.font.size = self.FONT_SIZE
        run.bold = True
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(12)
        pf.first_line_indent = Cm(0)

        p_toc = self.doc.add_paragraph()
        pf_toc = p_toc.paragraph_format
        pf_toc.space_before = Pt(0)
        pf_toc.space_after = Pt(0)
        pf_toc.first_line_indent = Cm(0)

        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run_begin = p_toc.add_run()
        run_begin._r.append(fld_begin)

        run_instr = p_toc.add_run()
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = r'TOC \o "1-3" \h \z \u'
        run_instr._r.append(instr)

        fld_separate = OxmlElement("w:fldChar")
        fld_separate.set(qn("w:fldCharType"), "separate")
        run_sep = p_toc.add_run()
        run_sep._r.append(fld_separate)

        run_placeholder = p_toc.add_run("[Actualice el índice con clic derecho > Actualizar campo]")
        run_placeholder.font.name = self.FONT
        run_placeholder.font.size = self.FONT_SIZE
        run_placeholder.font.color.rgb = RGBColor(128, 128, 128)
        run_placeholder.italic = True

        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run_end = p_toc.add_run()
        run_end._r.append(fld_end)

        self.doc.add_page_break()

    def add_title_page(self, title, authors, affiliation, course=None, instructor=None, date=None):
        self.add_page_number()
        lines = []
        lines.append(("\n" * 3))
        lines.append((title, True))
        lines.append("")
        lines.append((authors, False))
        lines.append((affiliation, False))
        if course:
            lines.append((course, False))
        if instructor:
            lines.append((instructor, False))
        if date:
            lines.append((date, False))

        for item in lines:
            if isinstance(item, str):
                if item == "":
                    p = self.doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run("")
                    run.font.size = self.FONT_SIZE
                elif item.startswith("\n"):
                    for _ in range(item.count("\n")):
                        p = self.doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        pf = p.paragraph_format
                        pf.space_before = Pt(0)
                        pf.space_after = Pt(0)
                continue
            text, bold = item
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            run = p.add_run(text)
            run.font.name = self.FONT
            run.font.size = self.FONT_SIZE
            run.bold = bold

        self.doc.add_page_break()

    def add_abstract(self, text, keywords=None):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Resumen")
        run.font.name = self.FONT
        run.font.size = self.FONT_SIZE
        run.bold = True
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.first_line_indent = Cm(0)

        p_body = self.doc.add_paragraph()
        run_body = p_body.add_run(text)
        run_body.font.name = self.FONT
        run_body.font.size = self.FONT_SIZE
        pf = p_body.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = self.LINE_SPACING
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.first_line_indent = Cm(0)

        if keywords:
            p_kw = self.doc.add_paragraph()
            run_label = p_kw.add_run("Palabras clave: ")
            run_label.font.name = self.FONT
            run_label.font.size = self.FONT_SIZE
            run_label.italic = True
            run_val = p_kw.add_run(", ".join(keywords))
            run_val.font.name = self.FONT
            run_val.font.size = self.FONT_SIZE
            pf = p_kw.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = self.LINE_SPACING
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.first_line_indent = Cm(0)

        self.doc.add_page_break()

    def add_heading_level1(self, text):
        return self._add_heading(text, "Heading 1")

    def add_heading_level2(self, text):
        return self._add_heading(text, "Heading 2")

    def add_heading_level3(self, text):
        return self._add_heading(text, "Heading 3")

    def add_heading_level4(self, text):
        p = self.doc.add_paragraph()
        run = p.add_run(text + ".")
        run.font.name = self.FONT
        run.font.size = self.FONT_SIZE
        run.bold = True
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = self.LINE_SPACING
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.first_line_indent = self.INDENT
        return p

    def add_heading_level5(self, text):
        p = self.doc.add_paragraph()
        run = p.add_run(text + ".")
        run.font.name = self.FONT
        run.font.size = self.FONT_SIZE
        run.bold = True
        run.italic = True
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = self.LINE_SPACING
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.first_line_indent = self.INDENT
        return p

    def _add_heading(self, text, style_name):
        p = self.doc.add_paragraph(style=style_name)
        run = p.add_run(text)
        run.font.name = self.FONT
        run.font.size = self.FONT_SIZE
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = self.LINE_SPACING
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.first_line_indent = Cm(0)
        self._set_keep_with_next(p)
        return p

    def add_paragraph(self, text):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = self.FONT
        run.font.size = self.FONT_SIZE
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = self.LINE_SPACING
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.first_line_indent = self.INDENT
        return p

    def add_block_quote(self, text):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = self.FONT
        run.font.size = self.FONT_SIZE
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = self.LINE_SPACING
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.first_line_indent = Cm(0)
        pf.left_indent = self.INDENT
        return p

    def add_reference(self, text):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = self.FONT
        run.font.size = self.FONT_SIZE
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = self.LINE_SPACING
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.first_line_indent = Inches(-0.5)
        pf.left_indent = Inches(0.5)
        return p

    def add_references_section(self, references):
        self.add_heading_level1("Referencias")
        for ref in references:
            self.add_reference(ref)

    def in_text_citation(self, author, year, page=None):
        text = f"{author} ({year})"
        if page:
            text = f"{author} ({year}, p. {page})"
        return text

    def parenthetical_citation(self, author, year, page=None):
        text = f"({author}, {year})"
        if page:
            text = f"({author}, {year}, p. {page})"
        return text

    def add_quote(self, text, author, year, page):
        p = self.doc.add_paragraph()
        run = p.add_run(f"{text} ")
        run.font.name = self.FONT
        run.font.size = self.FONT_SIZE
        run2 = p.add_run(self.parenthetical_citation(author, year, page))
        run2.font.name = self.FONT
        run2.font.size = self.FONT_SIZE
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = self.LINE_SPACING
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.first_line_indent = self.INDENT
        return p

    def _ref_paragraph(self):
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = self.LINE_SPACING
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.first_line_indent = Inches(-0.5)
        pf.left_indent = Inches(0.5)
        return p

    def _ref_run(self, p, text, italic=False):
        run = p.add_run(text)
        run.font.name = self.FONT
        run.font.size = self.FONT_SIZE
        run.italic = italic
        return run

    def add_journal_reference(self, authors, year, title, journal, volume, pages, doi=None, issue=None):
        p = self._ref_paragraph()
        self._ref_run(p, f"{authors} ({year}). ")
        self._ref_run(p, f"{title}. ", True)
        self._ref_run(p, f"{journal}, ")
        self._ref_run(p, f"{volume}", True)
        if issue:
            self._ref_run(p, f"({issue}), ")
        else:
            self._ref_run(p, ", ")
        self._ref_run(p, f"{pages}.")
        if doi:
            self._ref_run(p, f" https://doi.org/{doi}")
        return p

    def add_book_reference(self, authors, year, title, publisher, edition=None, doi=None):
        p = self._ref_paragraph()
        self._ref_run(p, f"{authors} ({year}). ")
        self._ref_run(p, f"{title}. ", True)
        if edition:
            self._ref_run(p, f"({edition} ed.). ")
        self._ref_run(p, f"{publisher}.")
        if doi:
            self._ref_run(p, f" https://doi.org/{doi}")
        return p

    def add_website_reference(self, authors, year, title, site_name, url, date_accessed=None):
        p = self._ref_paragraph()
        self._ref_run(p, f"{authors} ({year}). ")
        self._ref_run(p, f"{title}. ", True)
        self._ref_run(p, f"{site_name}. ")
        if date_accessed:
            self._ref_run(p, f"Recuperado el {date_accessed} de ")
        self._ref_run(p, url)
        return p

    def add_chapter_reference(self, authors, year, title, editors, book_title, pages, publisher, doi=None):
        p = self._ref_paragraph()
        self._ref_run(p, f"{authors} ({year}). ")
        self._ref_run(p, f"{title}. ", True)
        self._ref_run(p, f"En {editors} (Eds.), ")
        self._ref_run(p, f"{book_title} ", True)
        self._ref_run(p, f"(pp. {pages}). ")
        self._ref_run(p, f"{publisher}.")
        if doi:
            self._ref_run(p, f" https://doi.org/{doi}")
        return p

    def save(self, filename):
        self.add_page_number()
        self.doc.save(filename)
