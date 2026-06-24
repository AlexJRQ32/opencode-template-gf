import sys
sys.path.insert(0, r"{{USER_DIR}}\.agents\skills\apa7-docx\scripts")
from apa7 import APA7Doc
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy_header import copy_header_with_images

doc = APA7Doc()

def shade_cell(cell, color):
    el = OxmlElement("w:shd")
    el.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(el)

def header_row(table, texts, bg="1F3A5F"):
    row = table.rows[0]
    for i, t in enumerate(texts):
        c = row.cells[i]
        c.text = ""
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.5
        r = p.add_run(t)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.name = doc.FONT
        r.font.size = Pt(10)
        shade_cell(c, bg)

def style_data_cell(cell, bold=False):
    for p in cell.paragraphs:
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.5
        for r in p.runs:
            r.font.name = doc.FONT
            r.font.size = Pt(10)
            if bold:
                r.bold = True

def bullet(text):
    p = doc.doc.add_paragraph()
    run = p.add_run("\u2022 " + text)
    run.font.name = doc.FONT
    run.font.size = doc.FONT_SIZE
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = doc.LINE_SPACING
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = doc.INDENT

# ── TITLE PAGE ──
for _ in range(3):
    p = doc.doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

p = doc.doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
r = p.add_run("Propuesta de Proyecto Trabajo Comunal Universitario (TCU):")
r.bold = True
r.font.name = doc.FONT
r.font.size = doc.FONT_SIZE

p = doc.doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
r = p.add_run("Programaci\u00f3n Web B\u00e1sica: HTML, CSS y JavaScript")
r.bold = True
r.font.name = doc.FONT
r.font.size = doc.FONT_SIZE

for txt in [
    "",
    "{{STUDENT_NAME}}",
    "C\u00e9dula: 2-0874-0656",
    "Carrera de Ingenier\u00eda Inform\u00e1tica",
    "Universidad Hispanoamericana",
    "",
    "Sede Receptora:",
    "Centro Comunitario Inteligente (CECI) Micitt",
    "Junio, 2026",
]:
    p = doc.doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if txt:
        r = p.add_run(txt)
        r.font.name = doc.FONT
        r.font.size = doc.FONT_SIZE

doc.doc.add_page_break()

# ── TOC ──
doc.add_toc_page()

# ── TABLE 0: DATOS GENERALES ──
t0 = doc.doc.add_table(rows=6, cols=2)
t0.style = "Table Grid"
t0.alignment = WD_TABLE_ALIGNMENT.CENTER
header_row(t0, ["Datos Generales", "Datos de Carga Acad\u00e9mica"])

data0 = [
    (
        "Nombre del Proyecto: Programa Modular de Programaci\u00f3n Web B\u00e1sica",
        "Naturaleza del curso: Te\u00f3rico - Pr\u00e1ctico",
    ),
    (
        "Herramientas: VS Code, Google Chrome, GitHub Pages",
        "P\u00fablico Meta: Comunidad del CECI \u2014 personas sin conocimientos previos de programaci\u00f3n",
    ),
    (
        "Opci\u00f3n A: Entre semana - 2h/sesi\u00f3n, 2 sesiones/semana\nOpci\u00f3n B: Fin de semana - 4h/sesi\u00f3n, 1 sesi\u00f3n/semana\nModalidad: Presencial",
        "Horas TCU: 50h total / 40h curso / 10h preparaci\u00f3n",
    ),
    (
        "Requisitos: Conocimientos b\u00e1sicos en computaci\u00f3n",
        "Sede Receptora: CECI Micitt",
    ),
    (
        "Instructor: {{STUDENT_NAME}}",
        "Fecha de Inicio: Pendiente aprobaci\u00f3n (Estimado Julio 2026)",
    ),
]
for i, (a, b) in enumerate(data0, 1):
    t0.rows[i].cells[0].text = a
    t0.rows[i].cells[1].text = b
    style_data_cell(t0.rows[i].cells[0])
    style_data_cell(t0.rows[i].cells[1])

# ── 1. INTRODUCCIÃ“N ──
doc.add_heading_level1("Introducci\u00f3n a la Programaci\u00f3n")
doc.add_paragraph(
    "El mundo actual est\u00e1 impulsado por la tecnolog\u00eda, pero millones de personas siguen siendo solo "
    "consumidores pasivos de internet. Saben usar redes sociales, ver videos y buscar informaci\u00f3n, pero no "
    "entienden c\u00f3mo funciona la web ni c\u00f3mo podr\u00edan crear algo en ella."
)
doc.add_paragraph(
    "Este proyecto de TCU propone un curso introductorio de programaci\u00f3n web dirigido a personas de la "
    "comunidad del CECI que no saben absolutamente nada de computaci\u00f3n ni programaci\u00f3n. No se asume "
    "ning\u00fan conocimiento previo: ni qu\u00e9 es un archivo, una carpeta, un navegador, ni mucho menos c\u00f3digo."
)
doc.add_paragraph(
    "El curso cubre lo esencial de HTML, CSS y JavaScript en 10 semanas, llevando al estudiante desde prender "
    "una computadora hasta publicar su propia p\u00e1gina web personal en internet. El \u00e9nfasis est\u00e1 en la "
    "pr\u00e1ctica constante y en usar analog\u00edas del mundo real para evitar la jerga t\u00e9cnica."
)

# ── 2. JUSTIFICACIÃ“N ──
doc.add_heading_level1("Justificaci\u00f3n:")

doc.add_paragraph(
    "Como estudiante de Ingenier\u00eda Inform\u00e1tica, he podido observar que en nuestras comunidades existe "
    "una idea equivocada sobre lo que significa estar \u201cconectado\u201d. Muchas personas creen que saber usar "
    "una computadora para redes sociales y navegaci\u00f3n b\u00e1sica es suficiente, pero desconocen por completo "
    "c\u00f3mo funciona la tecnolog\u00eda que usan a diario."
)
doc.add_paragraph(
    "En las visitas a los centros CECI, es evidente que hay un gran potencial desperdiciado porque los usuarios "
    "se limitan a tareas b\u00e1sicas. Mi intenci\u00f3n con este TCU es romper esa barrera. Quiero que los "
    "participantes no solo consuman tecnolog\u00eda, sino que comiencen a crearla."
)
doc.add_paragraph(
    "Este proyecto no busca formar programadores expertos en diez semanas, sino despertar el pensamiento "
    "anal\u00edtico. Al aprender a estructurar una p\u00e1gina web desde cero, el ciudadano deja de ser un "
    "espectador pasivo de la tecnolog\u00eda y comienza a entender los principios que rigen el mundo digital "
    "que lo rodea."
)

# ── 3. OBJETIVOS ──
doc.add_heading_level1("Objetivos del Proyecto:")
doc.add_heading_level2("Objetivo General:")
doc.add_paragraph(
    "Desarrollar capacidades de pensamiento l\u00f3gico y anal\u00edtico en la comunidad del CECI mediante la "
    "ense\u00f1anza de fundamentos de la programaci\u00f3n web, con el fin de que logren estructurar, dise\u00f1ar "
    "y validar una p\u00e1gina web personal funcional utilizando HTML, CSS y JavaScript."
)
doc.add_heading_level2("Objetivos Espec\u00edficos:")
for b in [
    "Identificar los componentes b\u00e1sicos de una computadora, el navegador y el sistema de archivos para "
    "poder crear y guardar documentos digitales de forma independiente.",
    "Construir p\u00e1ginas web sencillas utilizando etiquetas HTML para estructurar contenido como t\u00edtulos, "
    "p\u00e1rrafos, im\u00e1genes, listas y enlaces.",
    "Aplicar estilos visuales b\u00e1sicos mediante CSS para modificar colores, fuentes y m\u00e1rgenes en una "
    "p\u00e1gina web.",
    "Incorporar interactividad m\u00ednima con JavaScript para responder a acciones del usuario como clics en "
    "botones.",
    "Publicar una p\u00e1gina web personal en internet utilizando GitHub Pages, permitiendo que sea accesible "
    "desde cualquier dispositivo.",
]:
    bullet(b)

# ── 4. CONTENIDO TEMÃ�TICO ──
doc.add_heading_level1("Contenido Tem\u00e1tico del Curso")

temas = [
    (
        "Tema 1: Fundamentos de Computaci\u00f3n B\u00e1sica.",
        [
            "Conceptos base: \u00bfQu\u00e9 es una computadora, hardware vs software?",
            "El sistema de archivos: \u00bfQu\u00e9 es un archivo, una carpeta, una unidad?",
            "El navegador web: \u00bfQu\u00e9 es, c\u00f3mo funciona, para qu\u00e9 sirve?",
            "Creaci\u00f3n del primer archivo: crear carpeta, guardar archivo .txt, abrirlo en el navegador",
        ],
    ),
    (
        "Tema 2: Introducci\u00f3n a HTML \u2014 Estructura de una P\u00e1gina Web.",
        [
            "\u00bfQu\u00e9 es HTML? Analog\u00eda del esqueleto humano",
            "Etiquetas b\u00e1sicas: <html>, <head>, <body>, <h1>, <p>",
            "Abrir VS Code, crear index.html, ver resultado en el navegador",
            "Pr\u00e1ctica: \u201cHola Mundo\u201d \u2014 primera p\u00e1gina web",
        ],
    ),
    (
        "Tema 3: Im\u00e1genes, Listas y Enlaces en HTML.",
        [
            "Insertar im\u00e1genes con <img>: atributo src y alt",
            "Crear enlaces con <a href>: conectar p\u00e1ginas y sitios externos",
            "Listas ordenadas y desordenadas: <ul>, <ol>, <li>",
            "Pr\u00e1ctica: P\u00e1gina \u201cMis cosas favoritas\u201d con lista y enlaces",
        ],
    ),
    (
        "Tema 4: Introducci\u00f3n a CSS \u2014 Apariencia Visual.",
        [
            "\u00bfQu\u00e9 es CSS? Analog\u00eda de la ropa y la apariencia",
            "Sintaxis b\u00e1sica: selector { propiedad: valor; }",
            "Colores de fondo y texto: background-color, color",
            "Pr\u00e1ctica: Aplicar colores a la p\u00e1gina personal",
        ],
    ),
    (
        "Tema 5: CSS \u2014 Flexbox para Layout.",
        [
            "\u00bfQu\u00e9 es Flexbox? Analog\u00eda de cajas flexibles que se ordenan solas",
            "Propiedades clave: display: flex, flex-direction, justify-content, align-items",
            "Crear una barra de navegaci\u00f3n horizontal con Flexbox",
            "Pr\u00e1ctica: Men\u00fa de navegaci\u00f3n funcional con enlaces",
        ],
    ),
    (
        "Tema 6: CSS \u2014 Grid B\u00e1sico y Personalizaci\u00f3n.",
        [
            "CSS Grid: display: grid, grid-template-rows: auto 1fr auto (layout completo)",
            "Diferencias con Flexbox: componentes vs. layout general",
            "Fuentes, bordes, sombras: font-family, border-radius, box-shadow",
            "Pr\u00e1ctica: Aplicar Grid al layout de la p\u00e1gina + personalizar estilos",
        ],
    ),
    (
        "Tema 7: Introducci\u00f3n a JavaScript \u2014 Interactividad B\u00e1sica.",
        [
            "\u00bfQu\u00e9 es JavaScript? Analog\u00eda de los movimientos del cuerpo",
            "Escribir en consola: console.log()",
            "Mostrar mensajes: alert()",
            "Pr\u00e1ctica: Bot\u00f3n que muestra un mensaje al hacerle clic",
        ],
    ),
    (
        "Tema 8: JavaScript \u2014 Eventos y Cambios en la P\u00e1gina.",
        [
            "Cambiar contenido con innerHTML y textContent",
            "Cambiar estilos desde JavaScript",
            "Evento onclick: responder a clics del usuario",
            "Pr\u00e1ctica: Bot\u00f3n que cambia el color de fondo al presionarlo",
        ],
    ),
]

for tit, items in temas:
    doc.add_heading_level2(tit)
    for it in items:
        bullet(it)

# ── 5. METODOLOGÃ�A ──
doc.add_heading_level1("Metodolog\u00eda de ense\u00f1anza:")
doc.add_paragraph(
    "La metodolog\u00eda de ense\u00f1anza est\u00e1 constituida para guiar al estudiante durante la clase como "
    "fuera de ella. Se utilizar\u00e1 un enfoque Te\u00f3rico-Pr\u00e1ctico distribuido de la siguiente manera:"
)
for b in [
    "Demostraciones guiadas: El instructor proyecta en vivo cada paso mientras los estudiantes replican en "
    "sus computadoras.",
    "Ejercicios pr\u00e1cticos en clase: Cada sesi\u00f3n termina con un ejercicio que el estudiante completa y "
    "ve funcionar en su navegador antes de irse.",
    "Analog\u00edas constantes: HTML es el esqueleto, CSS es la ropa, JavaScript son los movimientos. Nunca "
    "se usa jerga t\u00e9cnica sin explicarla primero.",
    "Repaso semanal: Los primeros 5-10 minutos de cada clase se repasa lo visto la semana anterior.",
    "Proyecto final integrador: A partir de la semana 7, los estudiantes trabajan en su p\u00e1gina personal, "
    "que se va construyendo clase a clase hasta la publicaci\u00f3n.",
]:
    bullet(b)

# ── 6. ESTRATEGIAS ──
doc.add_heading_level1("Estrategias de Aprendizaje:")
doc.add_paragraph(
    "Las estrategias que el estudiante adoptar\u00e1 para el cumplimiento de los objetivos son:"
)
for b in [
    "Ejercicios cortos en clase (Pr\u00e1cticas guiadas): Al final de cada sesi\u00f3n, completan un ejercicio "
    "individual.",
    "Mini-proyectos semanales: Cada semana construyen algo tangible que pueden mostrar.",
    "Trabajo activo en clase: Se fomenta la experimentaci\u00f3n \u2014 cambiar colores, probar etiquetas, "
    "romper y arreglar.",
    "Proyecto final progresivo: Desde la semana 7 construyen su p\u00e1gina personal, agregando capas cada clase.",
    "Auto-publicaci\u00f3n: En la \u00faltima clase, cada estudiante publica su p\u00e1gina en internet y la "
    "comparte con su familia.",
]:
    bullet(b)

# ── 7. RECURSOS ──
doc.add_heading_level1("Recursos did\u00e1cticos:")
for b in [
    "Equipo Tecnol\u00f3gico: PC o Laptop con sistemas operativos compatibles.",
    "Software Especializado: Instaladores de VS Code y Google Chrome.",
    "Acceso a internet: Para el intercambio de documentos de pr\u00e1ctica y recursos multimedia.",
    "Presentaci\u00f3n gr\u00e1fica: Uso de material digital para la explicaci\u00f3n de conceptos web.",
]:
    bullet(b)

# ── 8. EVALUACIÃ“N ──
doc.add_heading_level1("Evaluaci\u00f3n:")

t1 = doc.doc.add_table(rows=7, cols=2)
t1.style = "Table Grid"
t1.alignment = WD_TABLE_ALIGNMENT.CENTER
header_row(t1, ["Criterio de evaluaci\u00f3n", "Ponderaci\u00f3n"])

eval_items = [
    ("Ejercicios pr\u00e1cticos semanales (8 ejercicios)", "30%"),
    ("Proyecto final \u2014 P\u00e1gina web personal", "40%"),
    ("Participaci\u00f3n y trabajo en clase", "15%"),
    ("Mini-evaluaciones pr\u00e1cticas (quices)", "15%"),
]
for i, (c, pv) in enumerate(eval_items, 1):
    t1.rows[i].cells[0].text = c
    t1.rows[i].cells[1].text = pv
    style_data_cell(t1.rows[i].cells[0])
    style_data_cell(t1.rows[i].cells[1])
    t1.rows[i].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

r = t1.rows[5]
r.cells[0].text = "Total:"
r.cells[1].text = "100%"
for c in r.cells:
    style_data_cell(c, bold=True)
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.doc.add_paragraph()

# ── 9. CRONOGRAMA ──
doc.add_heading_level1("Cronograma del curso (10 semanas)")

t2 = doc.doc.add_table(rows=11, cols=4)
t2.style = "Table Grid"
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
header_row(t2, ["Semana", "Tema Principal", "Actividades de Aprendizaje", "Evaluaci\u00f3n"])

crono = [
    (
        "1",
        "Fundamentos de Computaci\u00f3n B\u00e1sica",
        "Conceptos de PC, archivos, carpetas, navegador. Crear carpeta, guardar .txt, abrir en navegador.",
        "Prueba Diagn\u00f3stica",
    ),
    ("2", "Mi primer HTML", "\u00bfQu\u00e9 es HTML? Etiquetas b\u00e1sicas, abrir VS Code, crear index.html.", "Pr\u00e1ctica Guiada"),
    ("3", "Im\u00e1genes, Listas y Enlaces", "Insertar im\u00e1genes, crear listas y enlaces. P\u00e1gina con foto y gustos.", "Quiz #1"),
    ("4", "CSS \u2014 Colores y Estilos", "Sintaxis CSS, colores de fondo y texto.", "Ejercicio semanal #1"),
    ("5", "CSS \u2014 Flexbox y Layout", "display: flex, justify-content, align-items. Crear barra de navegaci\u00f3n.", "Ejercicio semanal #2"),
    ("6", "CSS \u2014 Grid y Personalizaci\u00f3n", "Grid layout, fuentes, bordes, sombras. Layout completo de la p\u00e1gina.", "Pr\u00e1ctica en Clase"),
    ("7", "JavaScript \u2014 Primeros Pasos", "\u00bfQu\u00e9 es JS? alert(), console.log(), bot\u00f3n con mensaje.", "Quiz #2"),
    ("8", "JavaScript \u2014 Interactividad", "Cambiar contenido y estilos con eventos onclick.", "Ejercicio semanal #3"),
    ("9", "Proyecto Final", "Armar p\u00e1gina completa: foto, nombre, descripci\u00f3n, gustos, bot\u00f3n interactivo.", "Avance de Proyecto"),
    ("10", "Publicaci\u00f3n y Cierre", "Publicar en GitHub Pages. Compartir enlace. Cierre del curso.", "Proyecto Final"),
]
for i, (s, te, ac, ev) in enumerate(crono, 1):
    t2.rows[i].cells[0].text = s
    t2.rows[i].cells[1].text = te
    t2.rows[i].cells[2].text = ac
    t2.rows[i].cells[3].text = ev
    for c in t2.rows[i].cells:
        style_data_cell(c)

doc.doc.add_paragraph()

# Add schedule note
doc.add_paragraph(
    "Nota: El curso se ofrece en dos modalidades. Opción A: Entre semana, 2 horas por sesión, "
    "2 sesiones por semana (4 horas semanales). Opción B: Fin de semana, 4 horas por sesión, "
    "1 sesión por semana (4 horas semanales). El contenido y la duración total (40 horas de curso) "
    "son los mismos en ambas modalidades."
)

# ── 10. ANEXO DE RÃšBRICAS ──
doc.add_heading_level1("Anexo de rubricas:")

# 10a
doc.add_heading_level2("R\u00fabrica de Proyecto Programado")
t3 = doc.doc.add_table(rows=6, cols=4)
t3.style = "Table Grid"
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
header_row(t3, ["Criterio", "Excelente (5)", "Bueno (3)", "Insuficiente (1)"])

rub1 = [
    ("Funcionamiento", "Cumple 100% sin errores", "Funciona con detalles menores", "No funciona / Error de ejecuci\u00f3n"),
    ("Estructura HTML", "Uso \u00f3ptimo de etiquetas sem\u00e1nticas", "Uso funcional pero incompleto", "No utiliza HTML correctamente"),
    ("Estilo CSS", "Dise\u00f1o visual agradable, colores y m\u00e1rgenes", "Estilo b\u00e1sico funcional", "Sin CSS o muy m\u00ednimo"),
    ("Interactividad (JS)", "Bot\u00f3n funcional con respuesta al clic", "Bot\u00f3n presente pero no funciona", "Sin JavaScript"),
    ("Publicaci\u00f3n", "P\u00e1gina publicada en internet, enlace funcional", "Intent\u00f3 publicar", "No se public\u00f3"),
]
for i, (cri, exc, reg, ins) in enumerate(rub1, 1):
    t3.rows[i].cells[0].text = cri
    t3.rows[i].cells[1].text = exc
    t3.rows[i].cells[2].text = reg
    t3.rows[i].cells[3].text = ins
    for c in t3.rows[i].cells:
        style_data_cell(c)

doc.doc.add_paragraph()

# 10b
doc.add_heading_level2("R\u00fabrica de Tareas y Quices")
t4 = doc.doc.add_table(rows=5, cols=4)
t4.style = "Table Grid"
t4.alignment = WD_TABLE_ALIGNMENT.CENTER
header_row(t4, ["Aspecto a Evaluar", "Excelente (5)", "Bueno (3)", "Deficiente (1)"])

rub2 = [
    ("Puntualidad", "Entrega en hora y fecha", "Entrega tarde mismo d\u00eda", "No entreg\u00f3"),
    ("Funcionalidad", "Ejercicio funciona correctamente", "Errores menores", "No funciona"),
    ("Sintaxis HTML/CSS/JS", "Sin errores de escritura", "Errores leves", "Errores graves"),
    ("Independencia", "Completa sin ayuda", "Ayuda m\u00ednima", "Dependencia total"),
]
for i, (asp, exc, reg, ins) in enumerate(rub2, 1):
    t4.rows[i].cells[0].text = asp
    t4.rows[i].cells[1].text = exc
    t4.rows[i].cells[2].text = reg
    t4.rows[i].cells[3].text = ins
    for c in t4.rows[i].cells:
        style_data_cell(c)

doc.doc.add_paragraph()

# 10c
doc.add_heading_level2("R\u00fabrica de trabajo activo en clase")
t5 = doc.doc.add_table(rows=6, cols=4)
t5.style = "Table Grid"
t5.alignment = WD_TABLE_ALIGNMENT.CENTER
header_row(t5, ["Criterio de Actitud", "Excelente (2)", "Bueno (1)", "Deficiente (0)"])

rub3 = [
    ("Participaci\u00f3n Espont\u00e1nea", "Aporta soluciones siempre", "Participa si se le pide", "Nunca participa"),
    ("Pertinencia de Consultas", "Preguntas acordes al tema", "Preguntas fuera de tiempo", "No hace preguntas"),
    ("Respeto y Formalidad", "Lenguaje formal y respetuoso", "Lenguaje poco formal", "Falta de respeto"),
    ("Atenci\u00f3n en Clase", "Siempre atento y responde", "Distracci\u00f3n ocasional", "Ausente / No responde"),
    ("Preparaci\u00f3n Previa", "Conoce el tema al preguntar", "Conocimiento superficial", "No conoce el tema"),
]
for i, (cri, exc, reg, ins) in enumerate(rub3, 1):
    t5.rows[i].cells[0].text = cri
    t5.rows[i].cells[1].text = exc
    t5.rows[i].cells[2].text = reg
    t5.rows[i].cells[3].text = ins
    for c in t5.rows[i].cells:
        style_data_cell(c)

doc.doc.add_paragraph()

# ── REFERENCES ──
doc.add_references_section([
    "Mozilla Developer Network. (2024). HTML: Lenguaje de etiquetas de hipertexto. MDN Web Docs. https://developer.mozilla.org/es/docs/Web/HTML",
    "Mozilla Developer Network. (2024). CSS: Hojas de estilo en cascada. MDN Web Docs. https://developer.mozilla.org/es/docs/Web/CSS",
    "Mozilla Developer Network. (2024). JavaScript. MDN Web Docs. https://developer.mozilla.org/es/docs/Web/JavaScript",
    "Robbins, J. N. (2018). Learning Web Design: A Beginner\u2019s Guide to HTML, CSS, JavaScript, and Web Graphics (5th ed.). O\u2019Reilly Media.",
])

# ── SAVE ──
import os
temp = r"{{USER_DIR}}\AppData\Local\Temp\propuesta_web_tcu.docx"
doc.save(temp)
print("OK: Saved to " + temp)

# Also try to copy to final destination
final = r"{{USER_DIR}}\OneDrive\Documentos\Universidad-Hispanoamerica-Archivos\Archivos Independientes\TCU - Archivos\Propuesta Proyecto TCU - Programacion Web Basica.docx"
try:
    import shutil
    shutil.copy2(temp, final)
    print("OK: Also copied to " + final)

    # Copy header (with logo images) from the original CECI proposal
    header_source = r"{{USER_DIR}}\OneDrive\Documentos\Universidad-Hispanoamerica-Archivos\Archivos Independientes\TCU - Archivos\Propuesta Proyecto TCU - Introduccion a la Programacion.docx"
    copy_header_with_images(header_source, final)
except Exception as e:
    print("WARN: Could not copy to final: " + str(e))
    print("File is at: " + temp)


